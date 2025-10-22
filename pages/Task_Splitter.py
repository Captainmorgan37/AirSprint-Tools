from Home import password_gate
password_gate()
import re
from dataclasses import dataclass, field
from datetime import datetime, timedelta, date, time, timezone
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from collections import Counter, defaultdict
from collections.abc import Mapping
from typing import List, Dict, Any, Tuple, Optional, Set, Sequence

import pandas as pd
import pytz
from zoneinfo_compat import ZoneInfo
import streamlit as st
from pandas.api.types import is_scalar


from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from flight_leg_utils import (
    AIRPORT_TZ_FILENAME,
    FlightDataError,
    build_fl3xx_api_config,
    compute_departure_window_bounds,
    fetch_legs_dataframe,
    format_utc,
    is_customs_leg,
    load_airport_metadata_lookup,
    safe_parse_dt,
)

# ----------------------------
# App Config
# ----------------------------
st.set_page_config(page_title="Night-Shift Tail Splitter", layout="wide")
st.title("🛫 Night-Shift Tail Splitter")

st.caption(
    "Assign next-day tails to on-duty shifts as evenly as possible, while keeping all legs of a tail together."
)

UTC = timezone.utc
LOCAL_TZ = ZoneInfo("America/Edmonton")
DEPARTURE_WINDOW_START_UTC = time(hour=8, tzinfo=UTC)
DEPARTURE_WINDOW_END_UTC = time(hour=8, tzinfo=UTC)

# ----------------------------
# Types
# ----------------------------
@dataclass
class TailPackage:
    tail: str
    legs: int
    workload: float
    first_local_dt: datetime  # first dep local datetime for the day
    sample_legs: List[Dict[str, Any]]  # optional preview rows for UI (subset)
    has_priority: bool = False
    priority_labels: List[str] = field(default_factory=list)
    customs_legs: int = 0


# ----------------------------
# Helpers
# ----------------------------
_TAIL_BASE_WORKLOAD = 1.0
_BASE_LEG_WORKLOAD = 1.0
_CUSTOMS_LEG_BONUS = 0.25

def _to_local(dt: datetime, tz_name: str | None) -> datetime:
    if tz_name:
        try:
            return dt.astimezone(ZoneInfo(tz_name))
        except Exception:
            pass
    # Fallback: leave in original tz; if naive, assume UTC then convert to LOCAL_TZ so ordering is at least consistent
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=pytz.UTC)
    return dt.astimezone(LOCAL_TZ)


def _priority_label(value: Any) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, str):
        text = value.strip()
    else:
        text = str(value).strip()
    if not text:
        return None
    if "priority" in text.lower():
        return text
    return None


def _normalize_person_name(value: Any) -> str:
    if value is None:
        return ""
    if not is_scalar(value):
        # Lists/dicts sometimes store crew metadata; ignore them here.
        return ""
    try:
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        # Non-scalar objects or custom classes may raise here; treat them as empty.
        pass
    if not value:
        return ""
    text = str(value).strip()
    return text


def _member_display_name(member: Mapping[str, Any]) -> str:
    candidates = [
        member.get(key)
        for key in (
            "displayName",
            "display_name",
            "name",
            "fullName",
            "full_name",
        )
        if isinstance(member, Mapping)
    ]
    for candidate in candidates:
        name = _normalize_person_name(candidate)
        if name:
            return name
    if isinstance(member, Mapping):
        first = _normalize_person_name(
            member.get("firstName") or member.get("first_name")
        )
        last = _normalize_person_name(
            member.get("lastName") or member.get("last_name")
        )
        combined = " ".join(part for part in (first, last) if part)
        if combined.strip():
            return combined.strip()
    return ""


_PIC_KEYWORDS = {
    "pic",
    "picname",
    "captain",
    "pilotincommand",
    "pilot_in_command",
    "pilotcommand",
}

_SIC_KEYWORDS = {
    "sic",
    "sicname",
    "copilot",
    "firstofficer",
    "first_officer",
}


def _crew_names_from_row(row: Mapping[str, Any]) -> Tuple[str, str]:
    pic = ""
    sic = ""
    for key, value in row.items():
        if value is None:
            continue
        normalized_key = re.sub(r"[^a-z]", "", str(key).lower())
        if not normalized_key:
            continue
        name = _normalize_person_name(value)
        if not name:
            continue
        if not pic and normalized_key in _PIC_KEYWORDS:
            pic = name
        elif not sic and normalized_key in _SIC_KEYWORDS:
            sic = name
    crew_members = row.get("crewMembers")
    if isinstance(crew_members, list):
        for member in crew_members:
            if not isinstance(member, Mapping):
                continue
            role = str(member.get("role") or member.get("position") or "").lower()
            is_pic = bool(member.get("isPIC") or "pic" in role)
            is_sic = bool(member.get("isSIC") or "sic" in role or "first officer" in role)
            name = _member_display_name(member)
            if name:
                if not pic and is_pic:
                    pic = name
                elif not sic and is_sic:
                    sic = name
    return pic, sic


def _crew_names_from_package(pkg: "TailPackage") -> Tuple[str, str]:
    pic = ""
    sic = ""
    for leg in pkg.sample_legs:
        if isinstance(leg, Mapping):
            leg_pic, leg_sic = _crew_names_from_row(leg)
            if not pic and leg_pic:
                pic = leg_pic
            if not sic and leg_sic:
                sic = leg_sic
            if pic and sic:
                break
    return pic, sic


_TAIL_PLACEHOLDER_PREFIXES = ("ADD", "NEW", "TBD", "TEMP", "HOLD", "UNKNOWN", "UNK")
_TAIL_PLACEHOLDER_VALUES = {"", "NA", "N/A", "NONE", "NULL", "-"}
_TAIL_US_PATTERN = re.compile(r"^N[0-9]{1,5}[A-Z]{0,2}$")
_TAIL_HYPHEN_PATTERN = re.compile(r"^[A-Z0-9]{1,2}-[A-Z0-9]{2,5}$")
_TAIL_ALNUM_PATTERN = re.compile(r"^[A-Z0-9]{4,7}$")


def _is_valid_tail_registration(value: Any) -> bool:
    if not isinstance(value, str):
        return False
    candidate = value.strip().upper()
    if not candidate or candidate in _TAIL_PLACEHOLDER_VALUES:
        return False
    if any(ch.isspace() for ch in candidate):
        return False
    if candidate.startswith(_TAIL_PLACEHOLDER_PREFIXES):
        return False
    if len(candidate) < 3:
        return False
    if _TAIL_US_PATTERN.fullmatch(candidate):
        return True
    if _TAIL_HYPHEN_PATTERN.fullmatch(candidate):
        return True
    if "-" not in candidate and not any(ch.isdigit() for ch in candidate):
        return False
    if _TAIL_ALNUM_PATTERN.fullmatch(candidate):
        return True
    return False


def _default_target_date() -> date:
    """Return the default target date (two days ahead in local Mountain time)."""
    now_local = datetime.now(LOCAL_TZ)
    return (now_local + timedelta(days=2)).date()


def _default_shift_labels(count: int) -> List[str]:
    presets: Dict[int, List[str]] = {
        3: ["0500", "0800", "1200"],
        4: ["0500", "0600", "0800", "1200"],
        5: ["0500", "0600", "0800", "0900", "1200"],
    }
    if count in presets:
        return list(presets[count])
    return [f"Shift {i+1}" for i in range(count)]


def _alphabetical_suffix(index: int) -> str:
    if index < 0:
        return ""
    letters: List[str] = []
    while True:
        index, remainder = divmod(index, 26)
        letters.append(chr(ord("A") + remainder))
        if index == 0:
            break
        index -= 1
    return "".join(reversed(letters))


def _disambiguate_labels(labels: Sequence[str]) -> List[str]:
    normalized = [label or "Shift" for label in labels]
    counts = Counter(normalized)
    occurrences: defaultdict[str, int] = defaultdict(int)
    result: List[str] = []
    for base, original in zip(normalized, labels):
        total = counts[base]
        if total <= 1:
            result.append(original or base)
            continue
        occurrence = occurrences[base]
        occurrences[base] += 1
        suffix = _alphabetical_suffix(occurrence)
        result.append(f"{base}{suffix}")
    return result




# ----------------------------
# Data Fetch
# ----------------------------
@st.cache_data(show_spinner=False)
def fetch_next_day_legs(
    target_date: date,
    *,
    fl3xx_settings: Optional[Dict[str, Any]] = None,
    fetch_crew: bool = False,
) -> Tuple[pd.DataFrame, Dict[str, Any], Optional[Dict[str, Any]]]:
    """Fetch and normalise FL3XX legs for the tail splitter window."""

    try:
        config = build_fl3xx_api_config(fl3xx_settings)
    except FlightDataError as exc:
        st.error(str(exc))
        return pd.DataFrame(), {}, None
    except Exception as exc:  # pragma: no cover - defensive
        st.error(f"Error preparing FL3XX API configuration: {exc}")
        return pd.DataFrame(), {}, None

    window_start_utc, window_end_utc = compute_departure_window_bounds(
        target_date,
        start_time=DEPARTURE_WINDOW_START_UTC,
        end_time=DEPARTURE_WINDOW_END_UTC,
    )
    departure_window = (window_start_utc, window_end_utc)

    try:
        df, metadata, crew_summary = fetch_legs_dataframe(
            config,
            from_date=target_date,
            to_date=target_date + timedelta(days=2),
            departure_window=departure_window,
            fetch_crew=fetch_crew,
        )
    except Exception as exc:  # pragma: no cover - defensive
        st.error(f"Error fetching data from FL3XX API: {exc}")
        return pd.DataFrame(), {}, None

    normalization_stats = metadata.get("normalization_stats", {})
    skipped_subcharter = metadata.get("skipped_subcharter_legs") or normalization_stats.get(
        "skipped_subcharter", 0
    )
    if skipped_subcharter:
        st.info(
            "Skipped %d leg%s because the workflow contains 'Subcharter'."
            % (skipped_subcharter, "s" if skipped_subcharter != 1 else "")
        )

    legs_normalized = normalization_stats.get("legs_normalized", 0)
    window_meta = metadata.get("departure_window_utc") or {
        "start": format_utc(window_start_utc),
        "end": format_utc(window_end_utc),
    }
    window_counts = metadata.get("departure_window_counts", {})

    if df.empty:
        if legs_normalized == 0:
            if skipped_subcharter and skipped_subcharter == normalization_stats.get("candidate_legs", 0):
                return df, metadata, crew_summary
            st.warning("FL3XX API returned no recognizable legs for the selected date.")
        elif window_counts.get("within_window", 0) == 0:
            st.warning(
                "No FL3XX legs depart within the UTC window from %s to %s."
                % (window_meta.get("start"), window_meta.get("end"))
            )
        return df, metadata, crew_summary

    skipped_tail = normalization_stats.get("skipped_missing_tail", 0)
    skipped_time = normalization_stats.get("skipped_missing_dep_time", 0)
    if skipped_tail or skipped_time:
        skipped_total = skipped_tail + skipped_time
        st.warning(
            "Skipped %d leg%s missing required fields (tail missing: %d, departure time missing: %d)."
            % (
                skipped_total,
                "s" if skipped_total != 1 else "",
                skipped_tail,
                skipped_time,
            )
        )

    missing_tz_airports = metadata.get("missing_dep_tz_airports", [])
    tz_lookup_used = metadata.get("timezone_lookup_used", False)
    if missing_tz_airports:
        sample = ", ".join(missing_tz_airports)
        if len(sample) > 200:
            sample = sample[:197] + "..."
        message = (
            "Added timezone from airport lookup where possible. Update `%s` to cover: %s"
            % (AIRPORT_TZ_FILENAME, sample)
        )
        if tz_lookup_used:
            st.info(message)
        else:
            st.warning(
                "Unable to infer departure timezones automatically because `%s` was not found. "
                "Sample airports without tz: %s"
                % (AIRPORT_TZ_FILENAME, sample)
            )

    return df, metadata, crew_summary


def build_tail_packages(df: pd.DataFrame, target_date: date) -> Tuple[List[TailPackage], Set[str]]:
    if df.empty:
        return [], set()
    # Ensure required columns
    required = {"tail", "leg_id", "dep_time"}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing required columns in data: {missing}")

    df = df.copy()
    df["tail"] = df["tail"].astype(str)

    invalid_tails: Set[str] = set()

    def _valid_tail(value: Any) -> bool:
        tail_str = str(value)
        is_valid = _is_valid_tail_registration(tail_str)
        if not is_valid:
            invalid_tails.add(tail_str.strip())
        return is_valid

    df = df[df["tail"].map(_valid_tail)]
    if df.empty:
        return [], invalid_tails

    # Derive local first departure per tail for the day
    def first_local_for_tail(g: pd.DataFrame) -> datetime:
        # Filter legs that depart on target_date in their *local* timezone
        times_local: List[datetime] = []
        for _, row in g.iterrows():
            dt = safe_parse_dt(str(row["dep_time"]))
            tz_name = str(row.get("dep_tz", "")) or None
            dt_local = _to_local(dt, tz_name)
            if dt_local.date() == target_date:
                times_local.append(dt_local)
        if not times_local:
            # If none match exactly by local date, fall back to min local
            for _, row in g.iterrows():
                dt = safe_parse_dt(str(row["dep_time"]))
                tz_name = str(row.get("dep_tz", "")) or None
                times_local.append(_to_local(dt, tz_name))
        return min(times_local)

    airport_lookup = load_airport_metadata_lookup()

    packages: List[TailPackage] = []
    for tail, g in df.groupby("tail", sort=False):
        # Limit to target_date legs (by local date)
        legs_rows: List[Dict[str, Any]] = []
        all_rows: List[Dict[str, Any]] = []
        priority_values: Set[str] = set()
        for _, row in g.iterrows():
            row_dict = row.to_dict()
            is_customs = is_customs_leg(row_dict, airport_lookup)
            row_dict["is_customs_leg"] = is_customs
            all_rows.append(row_dict)
            priority_label = _priority_label(row_dict.get("workflowCustomName"))
            if priority_label:
                priority_values.add(priority_label)
            dt = safe_parse_dt(str(row_dict["dep_time"]))
            tz_name = str(row_dict.get("dep_tz", "")) or None
            dt_local = _to_local(dt, tz_name)
            if dt_local.date() == target_date:
                legs_rows.append(row_dict)
        # If none strictly on target_date by local, treat all as same-day package
        if not legs_rows:
            legs_rows = all_rows
        first_dt = first_local_for_tail(pd.DataFrame(legs_rows))
        customs_count = sum(1 for leg in legs_rows if leg.get("is_customs_leg"))
        workload = _TAIL_BASE_WORKLOAD
        for leg in legs_rows:
            workload += _BASE_LEG_WORKLOAD
            if leg.get("is_customs_leg"):
                workload += _CUSTOMS_LEG_BONUS
        packages.append(
            TailPackage(
                tail=str(tail),
                legs=len(legs_rows),
                workload=workload,
                first_local_dt=first_dt,
                sample_legs=legs_rows[:3],
                has_priority=bool(priority_values),
                priority_labels=sorted(priority_values),
                customs_legs=customs_count,
            )
        )
    return packages, invalid_tails


def assign_round_robin_by_first(packages: List[TailPackage], labels: List[str]) -> Dict[str, List[TailPackage]]:
    packages_sorted = sorted(packages, key=lambda p: p.first_local_dt)
    buckets: Dict[str, List[TailPackage]] = {lab: [] for lab in labels}
    for i, pkg in enumerate(packages_sorted):
        label = labels[i % len(labels)]
        buckets[label].append(pkg)
    return buckets


def assign_balanced_by_legs(packages: List[TailPackage], labels: List[str]) -> Dict[str, List[TailPackage]]:
    # Greedy bin-pack: biggest packages first → assign to bucket with lowest total legs
    buckets: Dict[str, List[TailPackage]] = {lab: [] for lab in labels}
    totals = {lab: 0.0 for lab in labels}

    def _workload(pkg: TailPackage) -> float:
        return pkg.workload if pkg.workload else float(pkg.legs)

    for pkg in sorted(packages, key=lambda p: _workload(p), reverse=True):
        # choose label with smallest total, then smallest count, then order
        label = sorted(
            labels,
            key=lambda lab: (totals[lab], len(buckets[lab]), labels.index(lab)),
        )[0]
        buckets[label].append(pkg)
        totals[label] += _workload(pkg)
    return buckets


def _offset_hours(dt: datetime) -> float:
    offset = dt.utcoffset()
    if offset is None:
        return 0.0
    return offset.total_seconds() / 3600


def assign_preference_weighted(
    packages: List[TailPackage],
    labels: List[str],
    label_weights: Optional[Sequence[float]] = None,
) -> Dict[str, List[TailPackage]]:
    if not packages or not labels:
        return {lab: [] for lab in labels}

    offsets = [_offset_hours(pkg.first_local_dt) for pkg in packages]
    min_off, max_off = min(offsets), max(offsets)
    span = max_off - min_off

    def _workload(pkg: TailPackage) -> float:
        return pkg.workload if pkg.workload else float(pkg.legs)

    total_workload = sum(_workload(pkg) for pkg in packages)

    weights: Dict[str, float] = {}
    for idx, lab in enumerate(labels):
        weight = 1.0
        if label_weights and idx < len(label_weights):
            try:
                weight = float(label_weights[idx])
            except (TypeError, ValueError):
                weight = 1.0
        if weight <= 0:
            weight = 1.0
        weights[lab] = weight

    total_weight = sum(weights.values()) or float(len(labels))
    baseline_target = total_workload / total_weight if total_weight else 0.0
    workload_targets = {lab: baseline_target * weights[lab] for lab in labels}

    # Use a tighter tolerance (10% of the even-share workload) so we aggressively
    # balance the workload while still respecting the east↔west preference ordering.
    tolerance = max(0.25, round(baseline_target * 0.10, 2)) if baseline_target else 0.25

    if len(labels) == 1:
        tz_targets = [max_off]
    elif span == 0:
        tz_targets = [max_off for _ in labels]
    else:
        step = span / (len(labels) - 1)
        tz_targets = [max_off - step * idx for idx in range(len(labels))]

    # Stage 1: assign each package to the shift that best matches its timezone
    # preference. We start by grouping everything strictly east→early and
    # west→late before considering workload balancing.
    buckets_by_index: List[List[TailPackage]] = [[] for _ in labels]
    totals_by_index: List[float] = [0.0 for _ in labels]
    preferred_index: Dict[str, int] = {}
    pkg_offsets: Dict[str, float] = {}

    for pkg in sorted(packages, key=lambda p: p.first_local_dt):
        pkg_offset = _offset_hours(pkg.first_local_dt)
        pkg_offsets[pkg.tail] = pkg_offset
        if len(labels) == 1 or span == 0:
            idx = 0
        else:
            relative = (max_off - pkg_offset) / span
            idx = int(round(relative * (len(labels) - 1)))
            idx = max(0, min(len(labels) - 1, idx))
        preferred_index[pkg.tail] = idx
        buckets_by_index[idx].append(pkg)
        totals_by_index[idx] += _workload(pkg)

    def _totals_delta(idx: int) -> float:
        label = labels[idx]
        return totals_by_index[idx] - workload_targets[label]

    max_iterations = len(packages) * max(1, len(labels) - 1) * 4
    iterations = 0

    # Stage 2: iteratively nudge packages forward/backward to balance workload
    # without letting them drift far from their preferred shift.
    while iterations < max_iterations:
        iterations += 1
        over_idx = max(range(len(labels)), key=_totals_delta)
        under_idx = min(range(len(labels)), key=_totals_delta)
        over_delta = _totals_delta(over_idx)
        under_delta = _totals_delta(under_idx)
        if over_delta <= tolerance and abs(under_delta) <= tolerance:
            break
        if over_delta <= 0 or under_delta >= 0:
            break

        step_direction = -1 if over_idx > under_idx else 1
        target_idx = over_idx + step_direction
        if target_idx < 0 or target_idx >= len(labels):
            break

        over_label = labels[over_idx]
        target_label = labels[target_idx]
        over_target_total = workload_targets[over_label]
        target_target_total = workload_targets[target_label]
        current_over_error = abs(totals_by_index[over_idx] - over_target_total)
        current_target_error = abs(totals_by_index[target_idx] - target_target_total)

        best_pkg: Optional[TailPackage] = None
        best_score: Optional[Tuple[float, float, float, float]] = None

        for pkg in buckets_by_index[over_idx]:
            work = _workload(pkg)
            new_over_total = totals_by_index[over_idx] - work
            new_target_total = totals_by_index[target_idx] + work
            new_over_error = abs(new_over_total - over_target_total)
            new_target_error = abs(new_target_total - target_target_total)
            delta_error = (new_over_error + new_target_error) - (
                current_over_error + current_target_error
            )
            pref_distance = abs(target_idx - preferred_index.get(pkg.tail, over_idx))
            tz_penalty = abs(pkg_offsets.get(pkg.tail, tz_targets[target_idx]) - tz_targets[target_idx])
            score = (delta_error, float(pref_distance), tz_penalty, work)
            if best_score is None or score < best_score:
                best_score = score
                best_pkg = pkg

        if not best_pkg or best_score is None:
            break

        delta_error = best_score[0]
        pref_distance = int(best_score[1])
        if delta_error > 0 and pref_distance > 1:
            # Moving this package would both worsen balance and ignore the
            # timezone preference; stop balancing.
            break
        if delta_error > 0 and pref_distance >= 1:
            # Try to find another option in the opposite direction if possible.
            alternative_idx = over_idx - step_direction
            if 0 <= alternative_idx < len(labels) and alternative_idx != target_idx:
                target_idx = alternative_idx
                # Restart loop to evaluate different direction.
                iterations -= 1
                continue
            break

        # Apply the move.
        buckets_by_index[over_idx].remove(best_pkg)
        buckets_by_index[target_idx].append(best_pkg)
        work = _workload(best_pkg)
        totals_by_index[over_idx] -= work
        totals_by_index[target_idx] += work

    result: Dict[str, List[TailPackage]] = {}
    for idx, label in enumerate(labels):
        pkgs = sorted(
            buckets_by_index[idx], key=lambda p: (p.first_local_dt, p.tail)
        )
        result[label] = pkgs
    return result


def buckets_to_df(
    buckets: Dict[str, List[TailPackage]],
    label_order: Optional[Sequence[str]] = None,
) -> pd.DataFrame:
    rows = []
    for label, pkgs in buckets.items():
        for pkg in sorted(pkgs, key=lambda p: (p.first_local_dt, p.tail)):
            rows.append({
                "Shift": label,
                "Tail": pkg.tail,
                "Legs": pkg.legs,
                "Customs Legs": pkg.customs_legs,
                "Workload": round(pkg.workload, 2),
                "First Local Dep": pkg.first_local_dt.strftime("%Y-%m-%d %H:%M %Z"),
                "Priority Flight": "Yes" if pkg.has_priority else "No",
                "Priority Detail": ", ".join(pkg.priority_labels) if pkg.priority_labels else "",
            })
    df = pd.DataFrame(rows)
    if not df.empty:
        categories: Optional[List[str]] = None
        if label_order:
            seen = list(pd.unique(df["Shift"]))
            ordered = [lab for lab in label_order if lab in seen]
            extras = [lab for lab in seen if lab not in ordered]
            categories = ordered + extras
        if categories:
            df["Shift"] = pd.Categorical(
                df["Shift"], categories=categories, ordered=True
            )
            df = df.sort_values(["Shift", "First Local Dep", "Tail"]).reset_index(
                drop=True
            )
            df["Shift"] = df["Shift"].astype(str)
        else:
            df = df.sort_values(["Shift", "First Local Dep", "Tail"]).reset_index(
                drop=True
            )
    return df


def summarize(
    df: pd.DataFrame, label_order: Optional[Sequence[str]] = None
) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    agg = (
        df.groupby("Shift", sort=False)
        .agg(
            Tails=("Tail", "count"),
            Legs=("Legs", "sum"),
            Customs=("Customs Legs", "sum"),
            Workload=("Workload", "sum"),
        )
        .reset_index()
    ).rename(columns={"Customs": "Customs Legs"})
    if label_order:
        seen = list(pd.unique(agg["Shift"]))
        ordered = [lab for lab in label_order if lab in seen]
        extras = [lab for lab in seen if lab not in ordered]
        categories = ordered + extras
        if categories:
            agg["Shift"] = pd.Categorical(
                agg["Shift"], categories=categories, ordered=True
            )
            agg = agg.sort_values("Shift").reset_index(drop=True)
            agg["Shift"] = agg["Shift"].astype(str)
    agg["Workload"] = agg["Workload"].round(2)
    # Add spread metrics
    total_workload = agg["Workload"].sum()
    total_shifts = agg.shape[0]
    target = total_workload / total_shifts if total_shifts else 0
    agg["Δ Workload vs Even"] = (agg["Workload"] - target).round(2)
    return agg


_DOCX_HEADERS = [
    "TAIL #",
    "CREW PIC",
    "CREW SIC",
    "FUEL",
    "CUSTOMS",
    "SLOT / PPR",
    "FLIGHT PLANS",
    "CREW BRIEF",
    "CONFIRMATION PIC",
    "CONFIRMATION SIC",
    "CHECK LIST",
    "RELEASE",
    "NOTES",
    "Priority Status",
]

_CHECKMARK = "✓"


def _apply_landscape(document: Document) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width


def _initialize_briefing_document(target_date: date) -> Document:
    document = Document()
    document.core_properties.title = f"{target_date} Shift Briefing"
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(8)

    title_para = document.add_paragraph(f"Daily Flight Sheet – {target_date}")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:
        title_run = title_para.runs[0]
    else:
        title_run = title_para.add_run(f"Daily Flight Sheet – {target_date}")
    title_run.font.size = Pt(16)
    title_run.bold = True

    _apply_landscape(document)
    return document


def _add_shift_table(
    document: Document,
    label: str,
    pkgs: List[TailPackage],
    priority_details: Dict[str, str],
) -> None:
    sorted_pkgs = sorted(pkgs, key=lambda p: (p.first_local_dt, p.tail))
    table_rows = len(sorted_pkgs) + 3  # header row + column headers + data + footer
    table = document.add_table(rows=table_rows, cols=len(_DOCX_HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Shift label header row spanning all columns
    top_cell = table.rows[0].cells[0]
    for merge_idx in range(1, len(_DOCX_HEADERS)):
        top_cell = top_cell.merge(table.rows[0].cells[merge_idx])
    top_paragraph = top_cell.paragraphs[0]
    top_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top_paragraph.add_run(label)
    run.bold = True
    run.font.size = Pt(14)

    # Column headers
    header_row = table.rows[1]
    for col_idx, header_text in enumerate(_DOCX_HEADERS):
        header_cell = header_row.cells[col_idx]
        header_cell.text = header_text
        header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        header_paragraph = header_cell.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.runs[0].font.bold = True

    # Data rows
    for row_offset, pkg in enumerate(sorted_pkgs):
        row = table.rows[row_offset + 2]
        pic_name, sic_name = _crew_names_from_package(pkg)
        values = [""] * len(_DOCX_HEADERS)
        values[0] = pkg.tail
        values[1] = pic_name
        values[2] = sic_name
        detail = priority_details.get(pkg.tail, "")
        cleaned_detail = ""
        if detail and not detail.lower().startswith("priority"):
            cleaned_detail = detail
        elif detail:
            cleaned_detail = detail.replace("priority", "", 1).strip() or detail
        if pkg.has_priority:
            values[13] = _CHECKMARK
            if cleaned_detail:
                values[13] = f"{values[13]} {cleaned_detail}".strip()
        elif cleaned_detail:
            values[13] = cleaned_detail
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_idx in {0, 13}:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer row for positioning/notes
    footer_row = table.rows[-1]
    positioning_cell = footer_row.cells[0]
    for merge_idx in range(1, max(1, len(_DOCX_HEADERS) // 2)):
        positioning_cell = positioning_cell.merge(footer_row.cells[merge_idx])
    positioning_cell.text = "POSITIONING:"
    positioning_cell.paragraphs[0].runs[0].bold = True

    notes_start = len(_DOCX_HEADERS) // 2
    notes_cell = footer_row.cells[notes_start]
    for merge_idx in range(notes_start + 1, len(_DOCX_HEADERS)):
        notes_cell = notes_cell.merge(footer_row.cells[merge_idx])
    notes_cell.text = "ADDITIONAL NOTES:"
    notes_cell.paragraphs[0].runs[0].bold = True


def _document_to_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _label_slug(label: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", label.strip()).strip("_")
    return slug or "shift"


def build_shift_briefing_docs(
    target_date: date,
    labels: List[str],
    buckets: Dict[str, List[TailPackage]],
    priority_details: Dict[str, str],
) -> tuple[bytes, Dict[str, bytes]]:
    combined_document = _initialize_briefing_document(target_date)

    for idx, label in enumerate(labels):
        pkgs = buckets.get(label, [])
        if idx > 0:
            combined_document.add_paragraph("")
        _add_shift_table(combined_document, label, pkgs, priority_details)

    combined_payload = _document_to_bytes(combined_document)

    per_shift_payloads: Dict[str, bytes] = {}
    for label in labels:
        shift_document = _initialize_briefing_document(target_date)
        _add_shift_table(shift_document, label, buckets.get(label, []), priority_details)
        per_shift_payloads[label] = _document_to_bytes(shift_document)

    return combined_payload, per_shift_payloads


# ----------------------------
# Sidebar: Inputs
# ----------------------------
st.sidebar.header("Inputs")

fl3xx_cfg: Dict[str, Any] = {}
try:
    if "fl3xx_api" in st.secrets:
        cfg = st.secrets["fl3xx_api"]
        if isinstance(cfg, Mapping):
            fl3xx_cfg = {str(k): cfg[k] for k in cfg}
        elif isinstance(cfg, dict):
            fl3xx_cfg = dict(cfg)
except Exception:
    # Accessing secrets outside Streamlit Cloud may raise; ignore gracefully.
    fl3xx_cfg = {}

has_live_credentials = bool(fl3xx_cfg.get("api_token") or fl3xx_cfg.get("auth_header"))

if not has_live_credentials:
    st.sidebar.info(
        "Add your FL3XX credentials to `.streamlit/secrets.toml` under `[fl3xx_api]` to enable live fetching.",
    )
else:
    st.sidebar.success("Using FL3XX credentials from Streamlit secrets.")

fetch_crew_default = bool(fl3xx_cfg.get("fetch_crew", True))
fetch_crew = st.sidebar.toggle(
    "Fetch crew details",
    value=fetch_crew_default and has_live_credentials,
    help="Retrieve crew information (PIC/SIC) for each flight. Requires additional API calls.",
    disabled=not has_live_credentials,
)

num_people = st.sidebar.number_input("Number of on-duty people", min_value=1, max_value=12, value=4, step=1)

default_labels = _default_shift_labels(int(num_people))
labels: List[str] = []
label_workloads: List[float] = []
for i in range(int(num_people)):
    lbl = st.sidebar.text_input(
        f"Label for person {i+1}",
        value=default_labels[i] if i < len(default_labels) else f"Shift {i+1}",
        key=f"label_{i}",
    )
    label_value = (lbl or "").strip()
    if not label_value:
        label_value = f"Shift {i+1}"
    labels.append(label_value)
    workload_percent = st.sidebar.slider(
        f"{label_value}: workload %",
        min_value=10,
        max_value=100,
        value=100,
        step=10,
        format="%d%%",
        key=f"workload_pct_{i}",
        help="Adjust this role's workload target as a percentage of a standard shift.",
    )
    label_workloads.append(workload_percent / 100.0)

raw_labels = list(labels)
labels = _disambiguate_labels(labels)
if labels != raw_labels:
    st.sidebar.caption("Duplicate shift names were suffixed automatically for clarity.")

# Date selection (default = two days ahead in local Mountain time)
selected_date = st.sidebar.date_input("Target date", value=_default_target_date())


# ----------------------------
# Main Action
# ----------------------------
fetch_col, reset_col = st.columns([4, 1])
with fetch_col:
    if st.button("🔄 Fetch & Assign", use_container_width=True):
        st.session_state["_run"] = True
with reset_col:
    if st.button("🧹 Clear cache", use_container_width=True):
        fetch_next_day_legs.clear()
        st.session_state.pop("_run", None)
        st.session_state["_cache_cleared"] = True

if st.session_state.pop("_cache_cleared", False):
    st.success("Cached data cleared. Fetch again to pull fresh data.")

# ----------------------------
# Processing & Output
# ----------------------------
if st.session_state.get("_run"):
    legs_df, _, crew_summary = fetch_next_day_legs(
        selected_date,
        fl3xx_settings=fl3xx_cfg if has_live_credentials else None,
        fetch_crew=bool(fetch_crew and has_live_credentials),
    )

    if legs_df.empty:
        st.warning("No legs returned for the selected date.")
        st.stop()

    if crew_summary and crew_summary.get("fetched"):
        st.sidebar.metric("Crew lookups", int(crew_summary["fetched"]))
        if crew_summary.get("errors"):
            st.sidebar.warning(f"Crew errors: {len(crew_summary['errors'])}")

    packages, invalid_tails = build_tail_packages(legs_df, selected_date)

    if invalid_tails:
        ignored = sorted(t for t in invalid_tails if t)
        if ignored:
            preview = ", ".join(ignored[:6])
            if len(ignored) > 6:
                preview += ", ..."
            st.info(
                "Ignored %d tail%s without an official registration: %s"
                % (
                    len(ignored),
                    "s" if len(ignored) != 1 else "",
                    preview,
                )
            )

    if not packages:
        st.info("No tail packages found for the selected date.")
        st.stop()

    priority_packages = [pkg for pkg in packages if pkg.has_priority]
    priority_tails = [pkg.tail for pkg in priority_packages]
    priority_details = {
        pkg.tail: ", ".join(pkg.priority_labels) if pkg.priority_labels else ""
        for pkg in priority_packages
    }

    st.subheader("Assignments")

    buckets = assign_preference_weighted(packages, labels, label_workloads)

    # Display per-shift tables
    tabs = st.tabs(labels)
    for i, lab in enumerate(labels):
        with tabs[i]:
            pkgs = buckets.get(lab, [])
            df = buckets_to_df({lab: pkgs}, label_order=[lab])
            if df.empty:
                st.write("No tails assigned.")
            else:
                st.dataframe(df, use_container_width=True, hide_index=True)
                total_legs = int(df["Legs"].sum())
                total_workload = round(float(df["Workload"].sum()), 2)
                customs_legs = int(df["Customs Legs"].sum())
                total_tails = int(df.shape[0])
                priority_total = int(sum(1 for p in pkgs if p.has_priority))
                col1, col2, col3, col4, col5 = st.columns(5)
                with col1:
                    st.metric("Total legs", total_legs)
                with col2:
                    st.metric("Workload-adjusted legs", total_workload)
                with col3:
                    st.metric("Customs legs", customs_legs)
                with col4:
                    st.metric("Tails", total_tails)
                with col5:
                    st.metric("Priority tails", priority_total)

    # Combined view
    combined_df = buckets_to_df(buckets, label_order=labels)

    if priority_tails:
        detail_list = [
            f"{tail} ({priority_details[tail]})" if priority_details[tail] else tail
            for tail in priority_tails
        ]
        st.warning(
            "Priority flights detected for: " + ", ".join(detail_list)
        )

    # Summary
    st.subheader("Summary")
    summary_df = summarize(combined_df, label_order=labels)
    st.dataframe(summary_df, use_container_width=True, hide_index=True)

    # Downloads
    doc_payload, per_shift_docs = build_shift_briefing_docs(
        selected_date, labels, buckets, priority_details
    )
    st.download_button(
        label="⬇️ Download daily flight sheet (DOCX)",
        data=doc_payload,
        file_name=f"daily_flight_sheet_{selected_date}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key="download-daily-flight-sheet",
    )

    if per_shift_docs:
        st.markdown("#### Individual shift documents")
        columns = st.columns(min(3, len(per_shift_docs)) or 1)
        for idx, label in enumerate(labels):
            payload = per_shift_docs.get(label)
            if not payload:
                continue
            column = columns[idx % len(columns)]
            with column:
                st.download_button(
                    label=f"⬇️ {label}",
                    data=payload,
                    file_name=f"daily_flight_sheet_{selected_date}_{_label_slug(label)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"download-shift-{idx}-{_label_slug(label)}",
                )
    st.download_button(
        label="⬇️ Download assignments (CSV)",
        data=combined_df.to_csv(index=False).encode("utf-8"),
        file_name=f"tail_assignments_{selected_date}.csv",
        mime="text/csv",
        use_container_width=True,
        key="download-assignments-csv",
    )

    st.success("Done. Adjust inputs and re-run as needed.")

